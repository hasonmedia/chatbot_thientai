import logging
from typing import Any, Dict, List
from pypdf import PdfReader
from docx import Document
import pandas as pd
from typing import Optional
import json
logger = logging.getLogger(__name__)
from datetime import datetime, date, time



async def extract_text_from_pdf(file_path: str) -> Optional[str]:
    try:
        reader = PdfReader(file_path)
        contents = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                contents.append(text.strip())

        return "\n\n".join(contents).strip()

    except Exception as e:
        logger.error(f"Lỗi đọc file PDF: {e}")
        return None

async def extract_text_from_docx(file_path: str) -> Optional[str]:

    try:
        doc = Document(file_path)
        contents = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                contents.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    contents.append(row_text)

        return "\n\n".join(contents).strip()

    except Exception as e:
        return None

async def extract_text_from_excel(file_path: str) -> Optional[str]:

    sheet_jsons = {}

    excel = pd.ExcelFile(file_path)
    for sheet in excel.sheet_names:
        try:
            df = pd.read_excel(file_path, sheet)
            if df.empty:
                continue
            
            
            rows_list = []
            for _, row in df.iterrows():
                row_dict = {col: val for col, val in row.items() if pd.notna(val)}
                if row_dict:
                    rows_list.append(row_dict)

            if rows_list:
                sheet_jsons[sheet] = rows_list

        except:
            continue
        
        final_str = "{\n"
        for sheet_name, rows in sheet_jsons.items():
            final_str += f'  "{sheet_name}": [\n'

            for row in rows:
                row_parts = []
                for k, v in row.items():
                    v_str = f'"{v}"' if isinstance(v, str) else str(v)
                    row_parts.append(f'"{k}": {v_str}')
                
                row_str = "{ " + ", ".join(row_parts) + " }"
                final_str += f"    {row_str},\n"

            final_str += "  ],\n"

        final_str += "}"

        return final_str


